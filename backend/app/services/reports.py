from __future__ import annotations

import json
from io import BytesIO
from typing import Any, Literal

from docx import Document
from docx.shared import Inches, Pt
from jinja2 import Environment, select_autoescape

from app.models.attempt import Attempt
from app.models.test import Test
from app.serializers import merge_config


def norm_interpretation(norm: int) -> str:
    if norm < 40:
        return "Низкий интерес"
    if norm < 55:
        return "Средний интерес"
    if norm < 75:
        return "Выше среднего"
    return "Выраженный интерес"


def enrich_metrics(metrics: list[dict]) -> list[dict]:
    out: list[dict] = []
    for m in metrics or []:
        name = str(m.get("name") or "Шкала")
        norm = int(m.get("value", 0))
        raw_val = m.get("raw")
        if raw_val is None:
            raw_f = float(norm)
        else:
            try:
                raw_f = float(raw_val)
            except (TypeError, ValueError):
                raw_f = float(norm)
        raw_display: int | float
        if abs(raw_f - round(raw_f)) < 1e-6:
            raw_display = int(round(raw_f))
        else:
            raw_display = round(raw_f, 2)
        desc = m.get("description")
        out.append(
            {
                "scale": name,
                "name": name,
                "raw": raw_f,
                "raw_display": raw_display,
                "norm": norm,
                "value": norm,
                "interpretation": norm_interpretation(norm),
                "description": (desc or "") if isinstance(desc, str) else "",
            }
        )
    return out


def metrics_to_lines(metrics: list[dict]) -> str:
    return "\n".join(
        f"• {m.get('name', '')}: {m.get('value', 0)}%"
        + (f" — {m['description']}" if m.get("description") else "")
        for m in (metrics or [])
    )


def _format_answer_value(q: dict, val: Any) -> str:
    qtype = q.get("type", "open")
    opts = {o["id"]: o.get("text", o["id"]) for o in (q.get("options") or [])}
    if qtype == "single":
        return str(opts.get(val, val) if val is not None else "—")
    if qtype == "multiple":
        if not isinstance(val, list):
            return str(val or "—")
        return ", ".join(str(opts.get(x, x)) for x in val) or "—"
    if qtype in ("scale", "number"):
        return str(val if val is not None else "—")
    if val is None:
        return "—"
    return str(val)


def _ascii_bars(enriched: list[dict], width: int = 20) -> str:
    lines: list[str] = []
    for m in enriched:
        norm = int(m["norm"])
        filled = int(round(norm / 100.0 * width))
        filled = max(0, min(width, filled))
        bar = "█" * filled + "░" * (width - filled)
        label = str(m["scale"])[:18]
        lines.append(f"{label:<18} {bar} {norm}%")
    return "\n".join(lines) if lines else "—"


def _client_recommendations(enriched: list[dict], leading: dict | None) -> list[str]:
    out: list[str] = []
    for m in sorted(enriched, key=lambda x: -x["norm"])[:4]:
        if m.get("description"):
            out.append(str(m["description"]))
    out.extend(
        [
            "Обратитесь к своему профориентологу для полной интерпретации результатов.",
            "Изучите программы обучения по рекомендованным направлениям.",
            "Пройдите дополнительные методики для уточнения профиля.",
        ]
    )
    return out[:8]


def _psychologist_recommendations(enriched: list[dict], interpretation: str) -> list[str]:
    recs: list[str] = []
    if len(enriched) >= 2:
        top = max(enriched, key=lambda x: x["norm"])
        second = sorted(enriched, key=lambda x: -x["norm"])[1]
        low = min(enriched, key=lambda x: x["norm"])
        recs.append(
            f"Обсудить с клиентом сочетание интересов в сферах «{top['scale']}» и «{second['scale']}»."
        )
        recs.append(
            f"Низкий показатель по «{low['scale']}» ({low['norm']}%) — учитывать при подборе направлений."
        )
    if interpretation.strip():
        for line in interpretation.strip().split("\n"):
            t = line.strip()
            if t and len(recs) < 6:
                recs.append(t)
    if not recs:
        recs.append("Сопоставьте профиль с методическими пособиями по данной методике.")
    return recs[:8]


def _metrics_chart_png(enriched: list[dict]) -> bytes | None:
    if not enriched:
        return None
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None
    try:
        labels = [str(m["scale"]) for m in enriched]
        values = [int(m["norm"]) for m in enriched]
        fig, ax = plt.subplots(figsize=(6.5, 3.2))
        ax.barh(labels[::-1], values[::-1], color="#2563eb")
        ax.set_xlim(0, 100)
        ax.set_xlabel("Нормализованный балл (0–100)")
        fig.tight_layout()
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


def build_report_context(attempt: Attempt, test: Test) -> dict[str, Any]:
    cd = attempt.client_data or {}
    c = merge_config(test.config) if test.config else merge_config({})
    qs = attempt.question_snapshot if attempt.question_snapshot else c.get("questions") or []
    by_id = {q.get("id", ""): q for q in qs if isinstance(q, dict) and q.get("id")}
    simple_metrics = (attempt.results or {}).get("simpleMetrics") if isinstance(attempt.results, dict) else None
    raw_metrics = simple_metrics or attempt.metrics or []
    enriched = enrich_metrics(raw_metrics)
    metric_breakdown: list[dict[str, Any]] = []
    if isinstance(simple_metrics, list):
        q_by_id = {str(q.get("id")): str(q.get("text") or q.get("id") or "") for q in qs if isinstance(q, dict)}
        for metric in simple_metrics:
            details = metric.get("details") if isinstance(metric, dict) else None
            if not isinstance(details, list) or not details:
                continue
            metric_breakdown.append(
                {
                    "name": str(metric.get("name") or "Метрика"),
                    "rows": [
                        {
                            "question": q_by_id.get(str(d.get("questionId")), str(d.get("questionId") or "")),
                            "score": d.get("score"),
                        }
                        for d in details
                    ],
                }
            )

    dt = attempt.completed_at or attempt.started_at
    completion_date = dt.strftime("%d.%m.%Y") if dt else ""
    completion_datetime = ""
    if dt:
        completion_datetime = dt.strftime("%d.%m.%Y %H:%M")
        if getattr(dt, "tzinfo", None):
            completion_datetime = f"{completion_datetime} {dt.tzname() or 'UTC'}"

    answers_list: list[dict[str, Any]] = []
    for ans in attempt.answers or []:
        qid = ans.get("questionId")
        q = by_id.get(qid, {})
        text_q = str(q.get("text") or qid or "")
        val = ans.get("value")
        answers_list.append(
            {
                "question": text_q,
                "answer": _format_answer_value(q, val),
                "questionId": qid,
            }
        )

    leading = max(enriched, key=lambda x: x["norm"]) if enriched else None
    interpretation = str(attempt.interpretation or "")
    test_version = str(c.get("demoVersion") if c.get("demoVersion") is not None else c.get("version") or "1")

    chart = {"labels": [m["scale"] for m in enriched], "values": [m["norm"] for m in enriched]}

    answers_raw = "\n".join(f"{a['question']}\n  → {a['answer']}" for a in answers_list)

    ctx: dict[str, Any] = {
        "clientName": str(cd.get("name", "")),
        "client_name": str(cd.get("name", "")),
        "clientEmail": str(cd.get("email", "") or ""),
        "client_email": str(cd.get("email", "") or ""),
        "clientAge": str(cd.get("age", "") or ""),
        "client_age": str(cd.get("age", "") or ""),
        "testTitle": test.title,
        "test_name": test.title,
        "session_id": str(attempt.id),
        "completion_date": completion_date,
        "completion_datetime": completion_datetime,
        "test_version": test_version,
        "interpretation": interpretation,
        "metrics": enriched,
        "metrics_ascii_bars": _ascii_bars(enriched),
        "metrics_chart_json": json.dumps(chart, ensure_ascii=False),
        "chart_labels": chart["labels"],
        "chart_values": chart["values"],
        "recommendations": _client_recommendations(enriched, leading),
        "psychologist_recommendations": _psychologist_recommendations(enriched, interpretation),
        "answers": answers_list,
        "answers_raw": answers_raw,
        "metrics_lines": metrics_to_lines(raw_metrics),
        "leading_scale": leading["scale"] if leading else "",
        "leading_norm": int(leading["norm"]) if leading else 0,
        "leading_description": str(leading.get("description") or "") if leading else "",
        "metrics_breakdown": metric_breakdown,
    }
    return ctx


_DEFAULT_CLIENT_HTML = """<!DOCTYPE html>
<html lang="ru">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{{ title }} — ПрофДНК</title>
  <style>
    body { font-family: system-ui, -apple-system, sans-serif; max-width: 42rem; margin: 2rem auto; padding: 0 1rem;
      line-height: 1.55; color: #1a1a1a; }
    h1 { font-size: 1.35rem; font-weight: 700; }
    h2 { font-size: 1.05rem; margin-top: 1.75rem; }
    table { border-collapse: collapse; width: 100%; font-size: 0.95rem; }
    th, td { border: 1px solid #e5e5e5; padding: 0.5rem 0.65rem; text-align: left; }
    th { background: #f8fafc; }
    .ascii { font-family: ui-monospace, monospace; white-space: pre-wrap; font-size: 0.82rem;
      background: #f8fafc; padding: 1rem; border-radius: 8px; overflow-x: auto; }
    ul { padding-left: 1.2rem; }
    .muted { color: #64748b; font-size: 0.85rem; margin-top: 2rem; }
  </style>
</head>
<body>
  <h1>Отчёт для клиента</h1>
  <h2>Личная информация</h2>
  <table>
    <tr><th>Поле</th><th>Значение</th></tr>
    <tr><td>Имя</td><td>{{ client_name or "—" }}</td></tr>
    <tr><td>Возраст</td><td>{{ client_age or "—" }}</td></tr>
    <tr><td>Дата прохождения</td><td>{{ completion_date or "—" }}</td></tr>
    <tr><td>Название теста</td><td>{{ test_name }}</td></tr>
  </table>

  <h2>Ваш профиль профессиональных интересов</h2>
  <p>Чем выше балл, тем больше направление соответствует вашим склонностям.</p>
  <p><strong>📊 График профиля</strong></p>
  <div class="ascii">{{ metrics_ascii_bars }}</div>

  {% if leading_scale %}
  <h2>🧭 Ведущая сфера</h2>
  <p><strong>{{ leading_scale }}</strong>{% if leading_norm %} ({{ leading_norm }}%){% endif %}
    {% if leading_description %} — {{ leading_description }}{% endif %}</p>
  {% endif %}

  <h2>Рекомендуемые направления</h2>
  <ul>
    {% for r in recommendations %}<li>{{ r }}</li>{% endfor %}
  </ul>

  <p class="muted">Отчёт сформирован автоматически на платформе ПрофДНК.</p>
</body>
</html>"""


_DEFAULT_PROFESSIONAL_HTML = """<!DOCTYPE html>
<html lang="ru">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{{ title }} — ПрофДНК</title>
  <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.1/dist/chart.umd.min.js"></script>
  <style>
    body { font-family: system-ui, -apple-system, sans-serif; max-width: 50rem; margin: 2rem auto; padding: 0 1rem;
      line-height: 1.5; color: #0f172a; }
    h1 { font-size: 1.3rem; }
    h2 { font-size: 1.05rem; margin-top: 1.5rem; }
    table { border-collapse: collapse; width: 100%; font-size: 0.9rem; }
    th, td { border: 1px solid #e2e8f0; padding: 0.45rem 0.6rem; text-align: left; }
    th { background: #f1f5f9; }
    .chart-wrap { max-width: 520px; margin: 1rem 0; }
    ul { padding-left: 1.2rem; }
    .muted { color: #64748b; font-size: 0.85rem; margin-top: 1.5rem; }
  </style>
</head>
<body>
  <h1>Профессиональный отчёт</h1>
  <h2>Данные клиента</h2>
  <table>
    <tr><th>Поле</th><th>Значение</th></tr>
    <tr><td>Имя</td><td>{{ client_name or "—" }}</td></tr>
    <tr><td>Возраст</td><td>{{ client_age or "—" }}</td></tr>
    <tr><td>ID прохождения</td><td>{{ session_id }}</td></tr>
    <tr><td>Дата и время</td><td>{{ completion_datetime or "—" }}</td></tr>
    <tr><td>Тест</td><td>{{ test_name }}</td></tr>
    <tr><td>Версия методики</td><td>{{ test_version }}</td></tr>
  </table>

  <h2>Сырые баллы и нормированные показатели</h2>
  <table>
    <tr><th>Шкала</th><th>Сырой балл</th><th>Нормализованный (0–100)</th><th>Интерпретация</th></tr>
    {% for m in metrics %}
    <tr>
      <td>{{ m.scale }}</td>
      <td>{{ m.raw_display }}</td>
      <td>{{ m.norm }}</td>
      <td>{{ m.interpretation }}</td>
    </tr>
    {% endfor %}
  </table>

  {% if metrics_breakdown %}
  <h2>Детализация метрик по вопросам</h2>
  {% for item in metrics_breakdown %}
    <p><strong>{{ item.name }}</strong></p>
    <table>
      <tr><th>Вопрос</th><th>Баллы</th></tr>
      {% for row in item.rows %}
      <tr><td>{{ row.question }}</td><td>{{ row.score }}</td></tr>
      {% endfor %}
    </table>
  {% endfor %}
  {% endif %}

  <h2>График профиля</h2>
  <div class="chart-wrap"><canvas id="profileChart" width="600" height="320"></canvas></div>
  <script>
    try {
      const labels = {{ chart_labels | tojson }};
      const values = {{ chart_values | tojson }};
      const ctx = document.getElementById('profileChart');
      if (ctx && labels && labels.length && typeof Chart !== 'undefined') {
        new Chart(ctx, {
          type: 'radar',
          data: {
            labels: labels,
            datasets: [{ label: 'Профиль (0–100)', data: values, fill: true,
              backgroundColor: 'rgba(37, 99, 235, 0.2)', borderColor: 'rgb(37, 99, 235)', pointBackgroundColor: 'rgb(37, 99, 235)' }]
          },
          options: {
            scales: { r: { beginAtZero: true, suggestedMax: 100 } }
          }
        });
      }
    } catch (e) { console.warn('Chart init:', e); }
  </script>

  {% if interpretation %}
  <h2>Интерпретация (авто)</h2>
  <p style="white-space:pre-wrap">{{ interpretation }}</p>
  {% endif %}

  <h2>Рекомендации для психолога</h2>
  <ul>{% for r in psychologist_recommendations %}<li>{{ r }}</li>{% endfor %}</ul>

  <h2>Ответы клиента (сводка)</h2>
  <table>
    <tr><th>Вопрос</th><th>Ответ</th></tr>
    {% for row in answers %}<tr><td>{{ row.question }}</td><td>{{ row.answer }}</td></tr>{% endfor %}
  </table>

  <p class="muted">Отчёт сформирован в реальном времени. Данные сохранены в системе.</p>
</body>
</html>"""


def build_report_html(
    title: str,
    template: str | None,
    context: dict[str, Any],
    kind: Literal["client", "professional"],
) -> str:
    default = _DEFAULT_CLIENT_HTML if kind == "client" else _DEFAULT_PROFESSIONAL_HTML
    tmpl_str = (template or "").strip() or default
    env = Environment(autoescape=select_autoescape(["html", "xml"]))
    env.filters["tojson"] = lambda v: json.dumps(v, ensure_ascii=False)
    tmpl = env.from_string(tmpl_str)
    merged = {**context, "title": title}
    return tmpl.render(**merged)


def _docx_add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(cell)


def _default_docx_client(doc: Document, title: str, ctx: dict[str, Any]) -> None:
    h = doc.add_heading(title, 0)
    for run in h.runs:
        run.font.size = Pt(18)
    doc.add_heading("Личная информация", level=1)
    _docx_add_table(
        doc,
        ["Поле", "Значение"],
        [
            ["Имя", ctx.get("client_name") or "—"],
            ["Возраст", ctx.get("client_age") or "—"],
            ["Дата прохождения", ctx.get("completion_date") or "—"],
            ["Название теста", ctx.get("test_name") or ""],
        ],
    )
    doc.add_heading("Профиль (текстовый график)", level=1)
    p = doc.add_paragraph(ctx.get("metrics_ascii_bars") or "—")
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    if ctx.get("leading_scale"):
        doc.add_heading("Ведущая сфера", level=1)
        ld = ctx.get("leading_description") or ""
        doc.add_paragraph(f"{ctx['leading_scale']} ({ctx.get('leading_norm', 0)}%). {ld}".strip())
    doc.add_heading("Рекомендации", level=1)
    for r in ctx.get("recommendations") or []:
        doc.add_paragraph(str(r), style="List Bullet")
    foot = doc.add_paragraph("Отчёт сформирован автоматически на платформе ПрофДНК.")
    for run in foot.runs:
        run.font.size = Pt(9)
        run.italic = True


def _default_docx_professional(doc: Document, title: str, ctx: dict[str, Any]) -> None:
    h = doc.add_heading(title, 0)
    for run in h.runs:
        run.font.size = Pt(18)
    doc.add_heading("Данные клиента", level=1)
    _docx_add_table(
        doc,
        ["Поле", "Значение"],
        [
            ["Имя", ctx.get("client_name") or "—"],
            ["Возраст", ctx.get("client_age") or "—"],
            ["ID прохождения", ctx.get("session_id") or ""],
            ["Дата и время", ctx.get("completion_datetime") or "—"],
            ["Тест", ctx.get("test_name") or ""],
            ["Версия методики", str(ctx.get("test_version") or "")],
        ],
    )
    doc.add_heading("Сырые и нормированные показатели", level=1)
    mrows = [
        [str(m["scale"]), str(m["raw_display"]), str(m["norm"]), str(m["interpretation"])]
        for m in (ctx.get("metrics") or [])
    ]
    _docx_add_table(doc, ["Шкала", "Сырой балл", "Норма 0–100", "Интерпретация"], mrows)

    if ctx.get("metrics_breakdown"):
        doc.add_heading("Детализация метрик по вопросам", level=1)
        for item in ctx.get("metrics_breakdown") or []:
            doc.add_paragraph(str(item.get("name", "Метрика")))
            rows = [[str(r.get("question", "")), str(r.get("score", ""))] for r in item.get("rows", [])]
            _docx_add_table(doc, ["Вопрос", "Баллы"], rows)

    png = _metrics_chart_png(ctx.get("metrics") or [])
    if png:
        doc.add_heading("График профиля", level=1)
        doc.add_picture(BytesIO(png), width=Inches(5.8))

    if ctx.get("interpretation"):
        doc.add_heading("Интерпретация", level=1)
        doc.add_paragraph(str(ctx["interpretation"]))

    doc.add_heading("Рекомендации для психолога", level=1)
    for r in ctx.get("psychologist_recommendations") or []:
        doc.add_paragraph(str(r), style="List Bullet")

    doc.add_heading("Ответы клиента", level=1)
    for row in ctx.get("answers") or []:
        doc.add_paragraph(f"{row.get('question', '')}")
        doc.add_paragraph(f"→ {row.get('answer', '')}", style="List Bullet 2")

    foot = doc.add_paragraph("Отчёт сформирован в реальном времени.")
    for run in foot.runs:
        run.font.size = Pt(9)
        run.italic = True


def build_report_docx(
    title: str,
    template: str | None,
    context: dict[str, Any],
    kind: Literal["client", "professional"],
) -> bytes:
    tmpl_str = (template or "").strip()
    if tmpl_str:
        env = Environment(autoescape=False)
        body = env.from_string(tmpl_str).render(**context, title=title)
        doc = Document()
        h = doc.add_heading(title, 0)
        for run in h.runs:
            run.font.size = Pt(18)
        for para in body.split("\n"):
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(11)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    doc = Document()
    if kind == "client":
        _default_docx_client(doc, title, context)
    else:
        _default_docx_professional(doc, title, context)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
